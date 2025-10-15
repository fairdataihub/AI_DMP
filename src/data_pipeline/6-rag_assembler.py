# rag_assembler.py
from pathlib import Path
from typing import Dict, Any, Optional
from logger import GLOBAL_LOGGER as log


class DMPAssembler:
    """Assemble DMP and save outputs."""

    @staticmethod
    def to_markdown(sections: Dict[str, str], project_info: Dict[str, Any]) -> str:
        title = f"DMP: {project_info.get('project_title', 'Research Project')}"
        lines = [f"# {title}", ""]
        if pi := project_info.get("pi_name"):
            lines.append(f"**PI:** {pi}")
        if inst := project_info.get("institution"):
            lines.append(f"**Institution:** {inst}")
        lines.append("")
        for key, txt in sections.items():
            lines.append(f"## {key.replace('_', ' ').title()}")
            lines.append(txt)
            lines.append("")
        return "\n".join(lines)

    @staticmethod
    def save_markdown(md_text: str, out_path: Path):
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(md_text, encoding="utf-8")
        return out_path

    @staticmethod
    def save_docx(md_text: str, out_path: Path):
        try:
            from docx import Document
            doc = Document()
            for line in md_text.splitlines():
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                else:
                    doc.add_paragraph(line)
            doc.save(str(out_path))
            return out_path
        except Exception as e:
            log.warning("Docx export failed", error=str(e))
            return None
